"""报告导出器 — 支持多种格式"""
import html
import os
import re
from pathlib import Path
from datetime import date
from typing import Optional

import config


# =============================================================================
# 格式检测
# =============================================================================

def detect_format(output_path: str) -> str:
    """根据文件扩展名检测目标格式"""
    ext = Path(output_path).suffix.lower()
    format_map = {
        ".md": "markdown",
        ".pdf": "pdf",
        ".docx": "word",
        ".html": "html",
        ".htm": "html",
        ".txt": "text",
    }
    return format_map.get(ext, "markdown")


# =============================================================================
# Markdown → HTML 转换（基础）
# =============================================================================

def markdown_to_html(markdown_text: str, title: str = "微信聊天记录报告") -> str:
    """将 Markdown 文本转换为 HTML"""
    import html

    # 简单的 Markdown → HTML 转换
    lines = markdown_text.split("\n")
    in_code_block = False
    in_table = False
    html_lines = []

    # 基本 CSS 样式
    css = """
    <style>
        body { font-family: -apple-system, "Microsoft YaHei", sans-serif;
               max-width: 900px; margin: 40px auto; padding: 0 20px;
               background: #fafafa; color: #333; line-height: 1.7; }
        h1 { color: #1a73e8; border-bottom: 2px solid #1a73e8; padding-bottom: 10px; }
        h2 { color: #333; border-left: 4px solid #1a73e8; padding-left: 12px; margin-top: 30px; }
        h3 { color: #555; margin-top: 20px; }
        code { background: #f0f0f0; padding: 2px 6px; border-radius: 3px;
               font-family: "SF Mono", Consolas, monospace; font-size: 0.9em; }
        pre { background: #f5f5f5; padding: 16px; border-radius: 8px;
              overflow-x: auto; border: 1px solid #e0e0e0; }
        pre code { background: none; padding: 0; }
        table { border-collapse: collapse; width: 100%; margin: 16px 0; }
        th { background: #1a73e8; color: white; padding: 10px 14px; text-align: left; }
        td { padding: 8px 14px; border-bottom: 1px solid #e0e0e0; }
        tr:nth-child(even) { background: #f9f9f9; }
        blockquote { border-left: 4px solid #ffa500; margin: 16px 0;
                    padding: 8px 16px; background: #fff8e1; color: #555; }
        strong { color: #d32f2f; }
        hr { border: none; border-top: 1px solid #e0e0e0; margin: 24px 0; }
        .header { background: white; padding: 20px 24px; border-radius: 12px;
                  box-shadow: 0 2px 8px rgba(0,0,0,0.08); margin-bottom: 30px; }
        .header h1 { margin: 0; font-size: 1.5em; color: #1a73e8; }
        .meta { color: #888; font-size: 0.9em; margin-top: 6px; }
    </style>
    """

    html_lines.append(f"<!DOCTYPE html>")
    html_lines.append(f"<html lang='zh-CN'>")
    html_lines.append(f"<head>")
    html_lines.append(f"<meta charset='utf-8'>")
    html_lines.append(f"<title>{html.escape(title)}</title>")
    html_lines.append(css)
    html_lines.append(f"</head>")
    html_lines.append(f"<body>")

    for line in lines:
        # 代码块
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                html_lines.append("<pre><code>")
            else:
                in_code_block = False
                html_lines.append("</code></pre>")
            continue
        if in_code_block:
            html_lines.append(html.escape(line))
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            html_lines.append(f"<h{level}>{html.escape(text)}</h{level}>")
            continue

        # 水平线
        if re.match(r"^[-*_]{3,}$", line.strip()):
            html_lines.append("<hr>")
            continue

        # 表格（简化处理）
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                html_lines.append("<table>")
            # 表头或数据行
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            tag = "th" if (len(html_lines) > 2 and "<table>" in html_lines[-2]) else "td"
            html_lines.append(f"<tr>" + "".join(f"<{tag}>{html.escape(c)}</{tag}>" for c in cells) + "</tr>")
            continue
        else:
            if in_table:
                in_table = False
                html_lines.append("</table>")

        # 有序/无序列表
        m = re.match(r"^(\s*)[-*]\s+(.+)", line)
        if m:
            indent = "  " * (len(m.group(1)) // 2)
            html_lines.append(f"{indent}<li>{_inline_format(m.group(2))}</li>")
            continue

        m = re.match(r"^(\s*)\d+\.\s+(.+)", line)
        if m:
            indent = "  " * (len(m.group(1)) // 2)
            html_lines.append(f"{indent}<li>{_inline_format(m.group(2))}</li>")
            continue

        # 引用
        if line.strip().startswith(">"):
            content = line.strip().lstrip("> ").strip()
            html_lines.append(f"<blockquote>{_inline_format(content)}</blockquote>")
            continue

        # 空行
        if not line.strip():
            html_lines.append("<br>")
            continue

        # 普通段落
        html_lines.append(f"<p>{_inline_format(line)}</p>")

    if in_table:
        html_lines.append("</table>")

    html_lines.append("</body></html>")
    return "\n".join(html_lines)


def _inline_format(text: str) -> str:
    """处理行内格式：粗体、斜体、代码"""
    text = html.escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"<a href='\2'>\1</a>", text)
    return text


# =============================================================================
# Markdown → Word (.docx)
# =============================================================================

def save_as_word(markdown_text: str, output_path: str, title: str = "微信聊天记录报告") -> None:
    """
    将 Markdown 保存为 Word 文档。
    需要安装 python-docx：pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError(
            "需要安装 python-docx：pip install python-docx\n"
            "安装命令：pip install python-docx"
        )

    doc = Document()

    # 标题
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = markdown_text.split("\n")
    in_code_block = False

    for line in lines:
        # 代码块
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                pass  # 代码块结束，下一行继续
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = "inQuote" if "inQuote" in doc.styles else "Normal"
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            doc.add_heading(text, level=min(level, 9))
            continue

        # 水平线
        if re.match(r"^[-*_]{3,}$", line.strip()):
            continue

        # 列表
        m = re.match(r"^(\s*)[-*]\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(m.group(2))
            continue

        m = re.match(r"^(\s*)\d+\.\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.add_run(m.group(2))
            continue

        # 引用
        if line.strip().startswith(">"):
            content = line.strip().lstrip("> ").strip()
            p = doc.add_paragraph(content)
            p.style = "Quote"
            continue

        # 空行
        if not line.strip():
            continue

        # 普通段落
        p = doc.add_paragraph(line)

    doc.save(output_path)


# =============================================================================
# Markdown → PDF（通过 HTML 中转）
# =============================================================================

def save_as_pdf(markdown_text: str, output_path: str, title: str = "微信聊天记录报告") -> None:
    """
    将 Markdown 保存为 PDF。
    支持两种后端（按优先级尝试）：
    1. weasyprint  pip install weasyprint
    2. pdfkit       pip install pdfkit (+ 系统安装 wkhtmltopdf)
    """
    html_content = markdown_to_html(markdown_text, title=title)

    # 尝试 weasyprint
    try:
        from weasyprint import HTML as WeasyHTML
        WeasyHTML(string=html_content).write_pdf(output_path)
        return
    except ImportError:
        pass

    # 尝试 pdfkit
    try:
        import pdfkit
        pdfkit.from_string(markdown_text, output_path, options={
            "encoding": "utf-8",
            "page-size": "A4",
            "margin-top": "20mm",
            "margin-bottom": "20mm",
            "margin-left": "15mm",
            "margin-right": "15mm",
        })
        return
    except ImportError:
        pass

    # 都不可用
    raise ImportError(
        "需要安装 PDF 生成库，二选一：\n"
        "  方式1（推荐）：pip install weasyprint\n"
        "  方式2：pip install pdfkit + 系统安装 wkhtmltopdf\n"
        "  下载 wkhtmltopdf：https://wkhtmltopdf.org/downloads.html"
    )


# =============================================================================
# Markdown → HTML（直接保存）
# =============================================================================

def save_as_html(markdown_text: str, output_path: str, title: str = "微信聊天记录报告") -> None:
    """将 Markdown 保存为 HTML 文件"""
    html_content = markdown_to_html(markdown_text, title=title)
    Path(output_path).write_text(html_content, encoding="utf-8")


# =============================================================================
# 统一导出入口
# =============================================================================

def export_report(
    report_text: str,
    output_path: str,
    title: str = "微信聊天记录报告",
) -> None:
    """
    将报告文本导出为指定格式。

    支持格式（根据扩展名自动判断）：
        .md   → Markdown（纯文本）
        .html → HTML（带样式）
        .pdf  → PDF（需安装 weasyprint 或 pdfkit）
        .docx → Word（需安装 python-docx）
        .txt  → 纯文本

    参数：
        report_text: 报告文本内容
        output_path: 输出文件路径（含扩展名）
        title: 文档标题（用于 HTML/PDF 元数据）
    """
    fmt = detect_format(output_path)

    if fmt == "markdown":
        Path(output_path).write_text(report_text, encoding="utf-8")

    elif fmt == "html":
        save_as_html(report_text, output_path, title=title)

    elif fmt == "pdf":
        save_as_pdf(report_text, output_path, title=title)

    elif fmt == "word":
        save_as_word(report_text, output_path, title=title)

    elif fmt == "text":
        # 去除 markdown 格式，保存为纯文本
        text = strip_markdown(report_text)
        Path(output_path).write_text(text, encoding="utf-8")

    else:
        raise ValueError(f"不支持的格式：{fmt}")


def strip_markdown(text: str) -> str:
    """去除 Markdown 格式，转为纯文本"""
    lines = text.split("\n")
    result = []
    for line in lines:
        # 去除标题标记
        line = re.sub(r"^#{1,6}\s+", "", line)
        # 去除粗体/斜体/代码标记
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"\*(.+?)\*", r"\1", line)
        line = re.sub(r"`(.+?)`", r"\1", line)
        # 去除链接
        line = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", line)
        # 去除列表标记
        line = re.sub(r"^\s*[-*+]\s+", "• ", line)
        line = re.sub(r"^\s*\d+\.\s+", "", line)
        # 去除引用
        line = re.sub(r"^\s*>\s*", "", line)
        result.append(line)
    return "\n".join(result)
